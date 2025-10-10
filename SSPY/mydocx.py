import copy
import os
import tempfile
import zipfile
from pathlib import Path
from sys import stderr

from docx.oxml.ns import qn
import xml.etree.ElementTree as ET


class DocxLoad:
    """解析docx文件"""

    def __init__(
        self,
        _path: str = None,
        parse_sheet: bool = True,
        parse_paragraphs: bool = False,
        if_print: bool = False):
        """
        Args:
            _path: 文件路径
        """
        self.__path = _path
        self.__if_print = if_print if isinstance(if_print, bool) else False
        self.__parse_sheet = parse_sheet if isinstance(parse_sheet, bool) else False
        self.__parse_paragraphs = parse_paragraphs if isinstance(parse_paragraphs, bool) else False
        if self.__parse_sheet == False and self.__parse_paragraphs == False: return
        if self.__if_print: print(_path)
        self.__sheets: list[list[list[str]]] = []
        self.__paragraphs: list[str] = []
        if self.__path is not None:
            self.__sheets, self.__paragraphs = self.parse_docx(self.__path)

    @staticmethod
    def repair_docx_for_docx(src_path: Path, if_print: bool = False) -> Path:
        with zipfile.ZipFile(src_path, 'r') as zin:
            # ① 真实存在集合
            real_names = set(zin.namelist())

            with tempfile.NamedTemporaryFile(suffix = '.docx', delete = False) as tmp:
                tmp_path = tmp.name

            with zipfile.ZipFile(tmp_path, 'w', compression = zipfile.ZIP_DEFLATED) as zout:
                for info in zin.infolist():
                    try:
                        data = zin.read(info)
                    except zipfile.BadZipFile:
                        if if_print:
                            print(f"[WARN] 跳过损坏 entry：{info.filename}")
                        continue

                    # ② 对 .rels 文件做“幽灵引用”清理
                    if info.filename.endswith('.rels'):
                        data = DocxLoad._sanitize_rels(data, info.filename, real_names, if_print)

                    zout.writestr(info, data)

        return Path(tmp_path)

    @staticmethod
    def _sanitize_rels(
        data: bytes,
        rels_name: str,
        real_names: set[str],
        if_print: bool = False) -> bytes:
        # 定义核心关系类型（只保留 key，不保留 URL 缩短）
        CORE_RELS = {
            'officeDocument',  # _rels/.rels -> word/document.xml
            'styles', 'fontTable', 'settings', 'theme', 'webSettings',  # word/_rels/document.xml.rels
            'footnotes', 'endnotes', 'header', 'footer',  # 可选但建议保留
        }
        try:
            root = ET.fromstring(data)
            namespace = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
            for rel in root.findall('r:Relationship', namespace):
                target = rel.get('Target')
                rel_type = rel.get('Type', '').split('/')[-1]  # 取最后一段
                if not target:
                    continue

                # ① 核心关系绝不删除
                if rel_type in CORE_RELS:
                    continue

                # ② 非核心关系才检查文件是否存在
                base = Path(rels_name).parent
                full = (base / target).resolve().as_posix()
                if full not in real_names:
                    if if_print:
                        print(f"[WARN] 删除幽灵引用：{rels_name} -> {target}")
                    root.remove(rel)
            return ET.tostring(root, encoding = 'utf-8', xml_declaration = True)
        except ET.ParseError:
            return data

    @staticmethod
    def get_merge_info(cell):
        """判断合并的cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 水平合并信息（跨列数）
        gridSpan = tcPr.find(qn('w:gridSpan'))
        horizontal_span = int(gridSpan.get(qn('w:val'))) if gridSpan is not None else 1

        # 垂直合并信息（是否为合并起始点）
        vMerge = tcPr.find(qn('w:vMerge'))
        is_vertical_start = False
        if vMerge is not None:
            v_merge_val = vMerge.get(qn('w:vMerge'))  # 注意这里修正了原代码的属性名错误
            is_vertical_start = (v_merge_val == 'restart')

        return horizontal_span > 1, horizontal_span, is_vertical_start

    def parse_docx(self, file_path: str):
        """
        Args:
            file_path:要解析的文件路径
        Returns:
            解析出来的内容
        """
        from docx import Document
        file_path = Path(file_path)
        tmp_path = None
        doc = None
        try:
            tmp_path = DocxLoad.repair_docx_for_docx(file_path, self.__if_print)

            doc = Document(tmp_path)
            all_tables: list[str] = []
            all_paragraphs: list[str] = []
            if self.__parse_sheet:
                for tbl_idx, table in enumerate(doc.tables):
                    try:
                        table_data = []
                        processed = set()

                        for row_idx, row in enumerate(table.rows):
                            try:
                                row_data = []
                                for col_idx, cell in enumerate(row.cells):
                                    if (row_idx, col_idx) in processed:
                                        continue
                                    row_data.append(cell.text.strip())
                                    # 处理水平合并 …
                                    is_h, span, is_v = DocxLoad.get_merge_info(cell)
                                    if is_h:
                                        for i in range(1, span):
                                            processed.add((row_idx, col_idx + i))
                                table_data.append(row_data)
                            except ValueError as e:
                                # ① 单行坏掉，只丢这一行
                                stderr.write(f"表格 {tbl_idx} 第 {row_idx} 行结构损坏，已跳过：{e}")
                                return None

                        all_tables.append(table_data)

                    except ValueError as e:
                        # ② 整个表格坏掉，直接丢
                        stderr.write(f"表格 {tbl_idx} 完全损坏，已跳过：{e}\n")
                        return None
            if self.__parse_paragraphs:
                for p in doc.paragraphs:
                    all_paragraphs.append(p.text.strip())
            return all_tables, all_paragraphs

        except ValueError as e:
            # 任何“不是 Word”的异常都抓
            if 'content type' in str(e).lower() or 'officeDocument' in str(e):
                # 再二次确认是不是 Excel
                with zipfile.ZipFile(tmp_path, 'r') as z:
                    ct = z.read('[Content_Types].xml').decode('utf-8', 'ignore').lower()
                    if 'spreadsheetml.main' in ct:
                        stderr.write(f'{file_path.name} 实为 Excel，已跳过\n')
                        return None  # 静默跳过
                return None

        finally:
            if doc is not None:
                del doc
            if tmp_path:
                tmp_path.unlink(missing_ok = True)

    @property
    def sheets(self) -> list[list[list]] | None:
        return copy.deepcopy(self.__sheets)

    @property
    def paragraphs(self) -> list[str] | None:
        return copy.deepcopy(self.__paragraphs)

    @property
    def sheets_without_enter(self) -> list[list[list]] | None:
        from .helperfunction import clean_enter
        if self.__sheets is None: return None
        outs = clean_enter(self.__sheets)
        return outs

    @property
    def path(self):
        return self.__path

    def get_sheet(self, index: int | str = None) -> list[list[str]] | None:
        """从文件中按照index内容读取一个表格"""
        if self.__sheets is None: return None
        from .fuzzy.search import searched_recursive as if_in
        if isinstance(index, int):
            if len(self.__sheets) > index:
                return copy.deepcopy(self.__sheets[index])
        if isinstance(index, str):  # 按照关键值查找sheet
            for sheet in self.__sheets:
                if if_in(index, sheet): return copy.deepcopy(sheet)
        return None

    def get_sheet_without_enter(self, index: int | str = None) -> list[list] | None:
        """从文件中按照index内容读取一个表格"""
        if self.__sheets is None: return None
        from .helperfunction import clean_enter
        sh = self.get_sheet(index = index)
        if sh is None:
            return None
        outs = clean_enter(sh)
        if isinstance(outs, list):
            return outs
        else:
            return None
